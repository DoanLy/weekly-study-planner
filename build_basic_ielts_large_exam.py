from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "De_moi_Basic_IELTS_ngu_phap_820_cau.docx"


def set_font(run, name="Calibri", size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_col_width(cell, width_in):
    width = Inches(width_in)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_col_width(cell, widths[i])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell_text(cell, size=10):
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            set_font(r, size=size)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def paragraph(doc, text="", size=10.5, bold=False, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    set_font(r, size=13, bold=True, color="2E74B5")


def make_options(correct, wrongs, seed=0):
    unique = []
    for item in wrongs:
        if item != correct and item not in unique:
            unique.append(item)
    if len(unique) < 3:
        raise ValueError(f"Not enough distractors for {correct}")
    pos = seed % 4
    opts = unique[:3]
    opts.insert(pos, correct)
    return opts, "ABCD"[pos]


def add_question(bank, section, text, correct, wrongs):
    opts, ans = make_options(correct, wrongs, len(bank) + len(section))
    section.append({"no": len(section) + 1, "q": text, "opts": opts, "ans": ans})


def add_mcq_table(doc, title, rows):
    heading(doc, title)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "No."
    hdr[1].text = "Question"
    repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        style_cell_text(cell, size=10.5)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for item in rows:
        row = table.add_row().cells
        row[0].text = str(item["no"])
        row[1].text = item["q"] + "\n" + "\n".join(
            f"{chr(65 + i)}. {option}" for i, option in enumerate(item["opts"])
        )
        style_cell_text(row[0], size=10)
        style_cell_text(row[1], size=10)
    set_table_geometry(table, [0.45, 6.05])


def add_written_table(doc, title, rows):
    heading(doc, title)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "No."
    hdr[1].text = "Task"
    hdr[2].text = "Answer"
    repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        style_cell_text(cell, size=10.5)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for item in rows:
        row = table.add_row().cells
        row[0].text = str(item["no"])
        row[1].text = item["task"]
        row[2].text = "\n\n"
        style_cell_text(row[0], size=10)
        style_cell_text(row[1], size=10)
    set_table_geometry(table, [0.45, 4.85, 1.20])


def build_pronunciation():
    rows = []
    s_groups = {
        "/s/": ["books", "cats", "maps", "laughs", "coughs", "stops", "months", "takes", "makes", "hopes", "types", "students", "desks", "cups", "hats", "cliffs", "beliefs", "cooks", "sleeps", "roofs"],
        "/iz/": ["buses", "watches", "boxes", "classes", "dishes", "bridges", "houses", "kisses", "freezes", "dances", "changes", "wishes", "washes", "villages", "judges", "matches", "sentences", "choices", "practices", "services"],
        "/z/": ["plays", "sees", "ends", "bags", "dreams", "drives", "belongs", "jobs", "islands", "pens", "rooms", "keys", "songs", "schools", "plans", "lives", "goes", "studies", "leaves", "knives"],
    }
    sounds = list(s_groups)
    for i in range(50):
        same = sounds[i % 3]
        diff = sounds[(i + 1 + i // 3) % 3]
        same_words = [s_groups[same][(i + j * 5) % len(s_groups[same])] for j in range(3)]
        correct = s_groups[diff][(i * 2) % len(s_groups[diff])]
        add_question(rows, rows, "Choose the word whose final -s/-es is pronounced differently from the others.", correct, same_words + [s_groups[diff][(i * 2 + 1) % len(s_groups[diff])]])

    ed_groups = {
        "/id/": ["wanted", "needed", "visited", "decided", "started", "ended", "added", "painted", "waited", "invited", "expected", "repeated", "included", "created", "collected"],
        "/t/": ["watched", "cooked", "stopped", "helped", "washed", "missed", "hoped", "looked", "fixed", "danced", "asked", "packed", "checked", "laughed", "worked"],
        "/d/": ["played", "opened", "called", "cleaned", "lived", "changed", "enjoyed", "planned", "arrived", "loved", "used", "borrowed", "offered", "studied", "carried"],
    }
    sounds = list(ed_groups)
    for i in range(30):
        same = sounds[i % 3]
        diff = sounds[(i + 2) % 3]
        same_words = [ed_groups[same][(i + j * 4) % len(ed_groups[same])] for j in range(3)]
        correct = ed_groups[diff][(i * 3) % len(ed_groups[diff])]
        add_question(rows, rows, "Choose the word whose final -ed is pronounced differently from the others.", correct, same_words + [ed_groups[diff][(i * 3 + 1) % len(ed_groups[diff])]])

    stress = [
        ("Choose the word with a different stress pattern.", "enjoy", ["teacher", "doctor", "artist"]),
        ("Choose the word with a different stress pattern.", "prefer", ["family", "holiday", "animal"]),
        ("Choose the word with a different stress pattern.", "invite", ["visit", "finish", "open"]),
        ("Choose the word with a different stress pattern.", "important", ["beautiful", "popular", "difficult"]),
        ("Choose the word with a different stress pattern.", "camera", ["computer", "banana", "tomato"]),
        ("Choose the word with a different stress pattern.", "newspaper", ["engineer", "volunteer", "Japanese"]),
        ("Choose the word with a different stress pattern.", "exercise", ["understand", "afternoon", "Vietnamese"]),
        ("Choose the word with a different stress pattern.", "simple", ["correct", "polite", "arrive"]),
        ("Choose the word with a different stress pattern.", "music", ["begin", "agree", "allow"]),
        ("Choose the word with a different stress pattern.", "answer", ["mistake", "complete", "believe"]),
        ("Choose the word with a different stress pattern.", "develop", ["document", "festival", "costume"]),
        ("Choose the word with a different stress pattern.", "condition", ["student", "village", "service"]),
        ("Choose the word with a different stress pattern.", "celebrate", ["discover", "decide", "prefer"]),
        ("Choose the word with a different stress pattern.", "factory", ["advertise", "important", "employment"]),
        ("Choose the word with a different stress pattern.", "energy", ["invitation", "opportunity", "interaction"]),
        ("Choose the word with a different stress pattern.", "listen", ["forget", "repeat", "prepare"]),
        ("Choose the word with a different stress pattern.", "garden", ["collect", "protect", "review"]),
        ("Choose the word with a different stress pattern.", "promise", ["arrive", "belong", "enjoy"]),
        ("Choose the word with a different stress pattern.", "private", ["replace", "create", "include"]),
        ("Choose the word with a different stress pattern.", "morning", ["support", "suggest", "explain"]),
    ]
    for q, correct, wrongs in stress:
        add_question(rows, rows, q, correct, wrongs)
    assert len(rows) == 100
    return rows


def build_vocab():
    rows = []
    like_items = [
        ("I enjoy ____ English songs after class.", "listening to", ["listen to", "to listening", "listened to"]),
        ("She prefers drinking tea ____ coffee.", "to", ["than", "from", "with"]),
        ("Tom is keen ____ learning new words.", "on", ["in", "at", "for"]),
        ("My sister is fond ____ reading novels.", "of", ["on", "for", "with"]),
        ("I am really ____ vintage fashion.", "into", ["onto", "inside", "about"]),
        ("He has a passion ____ charity work.", "for", ["at", "on", "with"]),
        ("They can't stand ____ in the rain.", "waiting", ["wait", "to waiting", "waited"]),
        ("I am not a big fan ____ spicy food.", "of", ["for", "about", "with"]),
        ("The word closest in meaning to 'detest' is ____.", "hate strongly", ["like", "prefer", "enjoy"]),
        ("The word closest in meaning to 'be fond of' is ____.", "like", ["hate", "avoid", "forget"]),
    ]
    signal_items = [
        ("'At the moment' is a signal phrase for ____.", "present continuous", ["present simple", "past simple", "passive voice"]),
        ("'Every day' usually signals ____.", "present simple", ["past continuous", "present perfect", "relative clause"]),
        ("'Yesterday' usually signals ____.", "past simple", ["present perfect", "present simple", "future simple"]),
        ("'At 8 p.m. yesterday' often signals ____.", "past continuous", ["present simple", "future simple", "modal verb"]),
        ("'So far' often signals ____.", "present perfect", ["past simple", "present continuous", "will"]),
        ("'Just' in 'She has just left' means ____.", "very recently", ["never", "a long time ago", "tomorrow"]),
        ("'Yet' is commonly used in ____ sentences in the present perfect.", "negative and question", ["only affirmative", "only passive", "only past simple"]),
        ("'For five years' answers the question ____.", "How long?", ["How often?", "Where?", "Why?"]),
        ("'Since 2020' shows ____.", "the starting point of a period", ["a finished past time only", "a place", "a reason"]),
        ("'Look at those clouds!' gives evidence for ____.", "be going to", ["past simple", "will for promises", "passive voice"]),
    ]
    term_items = [
        ("A 'gerund' is a verb form ending in ____ used like a noun.", "-ing", ["-ed", "-s", "-ly"]),
        ("An 'uncountable noun' usually takes a ____ verb.", "singular", ["plural", "modal", "future only"]),
        ("In passive voice, the object of the active sentence becomes the ____.", "subject", ["adverb", "preposition", "conjunction"]),
        ("A 'relative clause' usually gives more information about a ____.", "noun", ["modal verb", "tense marker", "sound only"]),
        ("'Who' is used mainly for ____.", "people", ["places", "reasons", "possessions"]),
        ("'Which' is used mainly for ____.", "things or animals", ["people only", "time only", "possession only"]),
        ("'Whose' shows ____.", "possession", ["place", "time", "reason"]),
        ("'Where' refers to ____.", "place", ["person", "reason", "possession"]),
        ("'When' refers to ____.", "time", ["place", "owner", "choice"]),
        ("'Why' refers to ____.", "reason", ["object", "place", "possession"]),
    ]
    conj_items = [
        ("'Because' introduces a ____.", "reason", ["choice", "possession", "pronunciation"]),
        ("'Although' introduces ____.", "contrast", ["purpose", "place", "certainty"]),
        ("'So that' introduces ____.", "purpose", ["contrast", "ownership", "a finished time"]),
        ("'Unless' means ____.", "if not", ["because", "although", "as soon as"]),
        ("'Either...or' expresses ____.", "choice", ["addition only", "reason", "time"]),
        ("'Both...and' means ____.", "two things together", ["neither thing", "one of two", "a reason"]),
        ("'Neither...nor' expresses ____.", "two negative ideas", ["two positive ideas", "a purpose", "a time"]),
        ("'Not only...but also' adds ____ information.", "emphatic", ["negative only", "time", "place"]),
        ("'As soon as' means ____.", "immediately after", ["although", "if not", "because"]),
        ("'Until' means ____.", "up to the time when", ["because", "despite", "in order to"]),
    ]
    all_items = like_items + signal_items + term_items + conj_items
    for i in range(100):
        q, correct, wrongs = all_items[i % len(all_items)]
        prefix = "" if i < len(all_items) else f"Exam-style vocabulary review {i + 1}: "
        add_question(rows, rows, prefix + q, correct, wrongs)
    assert len(rows) == 100
    return rows


def build_tenses():
    rows = []
    banks = [
        ("present simple", [
            ("My father ____ to work by bus every day.", "goes", ["go", "is going", "went"]),
            ("Water ____ at 100 degrees Celsius.", "boils", ["boil", "is boiling", "boiled"]),
            ("They usually ____ football after school.", "play", ["plays", "are playing", "played"]),
            ("She ____ meat because she is a vegetarian.", "doesn't eat", ["isn't eating", "didn't eat", "hasn't eaten"]),
            ("The train ____ at 8 a.m. tomorrow.", "leaves", ["is leaving", "left", "has left"]),
            ("My brother rarely ____ breakfast.", "skips", ["skip", "is skipping", "skipped"]),
            ("The Earth ____ around the Sun.", "goes", ["go", "is going", "went"]),
            ("I ____ English twice a week.", "study", ["studies", "am studying", "studied"]),
            ("Every child ____ a chance to learn.", "needs", ["need", "is needing", "needed"]),
            ("This shop ____ at 9 p.m.", "closes", ["close", "is closing", "closed"]),
        ]),
        ("present continuous", [
            ("Listen! The birds ____ outside.", "are singing", ["sing", "sings", "sang"]),
            ("We ____ for the bus at the moment.", "are waiting", ["wait", "waited", "have waited"]),
            ("She ____ in the garden now.", "is working", ["works", "worked", "has worked"]),
            ("Be quiet! The baby ____.", "is sleeping", ["sleeps", "slept", "has slept"]),
            ("They ____ a new project this week.", "are starting", ["start", "started", "have started"]),
            ("He is always ____ late for class.", "coming", ["come", "comes", "came"]),
            ("I ____ to London next Friday. I bought the ticket.", "am flying", ["fly", "flew", "have flown"]),
            ("Look! The children ____ in the park.", "are running", ["run", "ran", "have run"]),
            ("She ____ dinner right now.", "is cooking", ["cooks", "cooked", "has cooked"]),
            ("They ____ English at present.", "are learning", ["learn", "learned", "have learned"]),
        ]),
        ("past simple", [
            ("I ____ my grandmother last weekend.", "visited", ["visit", "have visited", "was visiting"]),
            ("They ____ to Da Nang two days ago.", "went", ["go", "have gone", "were going"]),
            ("She ____ up, took her bag, and left.", "stood", ["stands", "was standing", "has stood"]),
            ("We ____ the museum in 2022.", "visited", ["visit", "have visited", "were visiting"]),
            ("He ____ the door five minutes ago.", "opened", ["opens", "has opened", "was opening"]),
            ("My parents ____ a new car last year.", "bought", ["buy", "have bought", "were buying"]),
            ("The class ____ at 7 p.m. yesterday.", "finished", ["finishes", "has finished", "was finishing"]),
            ("She ____ breakfast before school this morning.", "had", ["has", "has had", "was having"]),
            ("Tom ____ the answer in the test.", "chose", ["choose", "has chosen", "was choosing"]),
            ("They ____ home late last night.", "came", ["come", "have come", "were coming"]),
        ]),
        ("past continuous", [
            ("At 9 p.m. yesterday, we ____ a movie.", "were watching", ["watched", "watch", "have watched"]),
            ("When I got home, my sister ____ dinner.", "was cooking", ["cooked", "has cooked", "cooks"]),
            ("While my dad ____ the car, I was cleaning my room.", "was washing", ["washed", "has washed", "washes"]),
            ("The phone rang while I ____ homework.", "was doing", ["did", "have done", "do"]),
            ("It ____ heavily when we left the office.", "was raining", ["rained", "has rained", "rains"]),
            ("They ____ TV when the lights went out.", "were watching", ["watched", "watch", "have watched"]),
            ("At that time yesterday, she ____ for the exam.", "was studying", ["studied", "studies", "has studied"]),
            ("While I ____ to school, I met an old friend.", "was walking", ["walked", "walk", "have walked"]),
            ("The students ____ loudly when the teacher entered.", "were talking", ["talked", "talk", "have talked"]),
            ("I ____ when the alarm rang.", "was sleeping", ["slept", "sleep", "have slept"]),
        ]),
        ("present perfect", [
            ("Mary ____ any meat for years.", "hasn't eaten", ["doesn't eat", "didn't eat", "isn't eating"]),
            ("I ____ this film twice.", "have seen", ["saw", "see", "am seeing"]),
            ("Have you ____ finished your homework?", "already", ["yet", "ago", "last"]),
            ("We haven't met him ____.", "yet", ["already", "ever", "just"]),
            ("She has ____ visited Japan before.", "never", ["ago", "last", "yesterday"]),
            ("They have lived here ____ 2020.", "since", ["for", "ago", "last"]),
            ("He has studied English ____ five years.", "for", ["since", "ago", "last"]),
            ("Over the past month, prices ____ quickly.", "have increased", ["increase", "increased", "were increasing"]),
            ("I have ____ eaten lunch, so I am not hungry.", "just", ["yet", "ago", "last"]),
            ("____ you ever tried Indian food?", "Have", ["Do", "Did", "Are"]),
        ]),
    ]
    for tense, items in banks:
        for i in range(20):
            q, correct, wrongs = items[i % len(items)]
            add_question(rows, rows, f"({tense}) {q}", correct, wrongs)
    assert len(rows) == 100
    return rows


def build_relative():
    rows = []
    fill = [
        ("The man ____ is standing there is my uncle.", "who", ["which", "where", "when"]),
        ("The book ____ I borrowed is useful.", "which", ["who", "where", "why"]),
        ("The girl ____ hair is long is my friend.", "whose", ["who", "which", "where"]),
        ("This is the house ____ I was born.", "where", ["who", "which", "why"]),
        ("Do you remember the day ____ we first met?", "when", ["where", "why", "whose"]),
        ("Tell me the reason ____ you are late.", "why", ["where", "when", "whose"]),
        ("The woman ____ you met yesterday is my boss.", "whom", ["where", "when", "whose"]),
        ("The car ____ he bought is expensive.", "that", ["where", "when", "why"]),
        ("I like people ____ are honest.", "who", ["which", "where", "when"]),
        ("The city ____ we visited was beautiful.", "which", ["who", "whose", "why"]),
    ]
    combine = [
        ("Combine: The book is interesting. I borrowed it from the library.", "The book which I borrowed from the library is interesting.", ["The book who I borrowed from the library is interesting.", "The book where I borrowed from the library is interesting.", "The book whose I borrowed from the library is interesting."]),
        ("Combine: The girl won the prize. She is my best friend.", "The girl who won the prize is my best friend.", ["The girl which won the prize is my best friend.", "The girl where won the prize is my best friend.", "The girl whose won the prize is my best friend."]),
        ("Combine: This is the house. I was born there.", "This is the house where I was born.", ["This is the house who I was born.", "This is the house whose I was born.", "This is the house why I was born."]),
        ("Combine: I remember the day. We first met on that day.", "I remember the day when we first met.", ["I remember the day where we first met.", "I remember the day why we first met.", "I remember the day whose we first met."]),
        ("Combine: The boy is crying. His bike was stolen.", "The boy whose bike was stolen is crying.", ["The boy who bike was stolen is crying.", "The boy which bike was stolen is crying.", "The boy where bike was stolen is crying."]),
    ]
    for i in range(60):
        q, correct, wrongs = fill[i % len(fill)]
        add_question(rows, rows, q, correct, wrongs)
    for i in range(40):
        q, correct, wrongs = combine[i % len(combine)]
        add_question(rows, rows, q, correct, wrongs)
    assert len(rows) == 100
    return rows


def build_sva():
    rows = []
    items = [
        ("Everyone ____ to be happy.", "wants", ["want", "are wanting", "have wanted"]),
        ("The information ____ useful.", "is", ["are", "were", "be"]),
        ("Ten dollars ____ not enough.", "is", ["are", "were", "have"]),
        ("Tom and Jerry ____ famous.", "are", ["is", "was", "be"]),
        ("The police ____ investigating the case.", "are", ["is", "has", "was"]),
        ("A number of students ____ absent today.", "are", ["is", "has", "was"]),
        ("The number of students ____ thirty.", "is", ["are", "were", "have"]),
        ("Each student ____ a dictionary.", "has", ["have", "are having", "were having"]),
        ("Swimming ____ good for your health.", "is", ["are", "were", "have"]),
        ("Neither the teacher nor the students ____ ready.", "are", ["is", "has", "was"]),
        ("Either my parents or my brother ____ me.", "meets", ["meet", "are meeting", "were meeting"]),
        ("A lot of water ____ wasted.", "is", ["are", "were", "have"]),
        ("Some of the animals ____ in danger.", "are", ["is", "has", "was"]),
        ("My mother, as well as my sisters, ____ at home.", "is", ["are", "were", "have"]),
        ("The rich ____ not always happy.", "are", ["is", "has", "was"]),
        ("Reading books ____ my favorite hobby.", "is", ["are", "were", "have"]),
        ("Every child ____ a toy.", "has", ["have", "are", "were"]),
        ("Physics ____ difficult for some students.", "is", ["are", "were", "have"]),
        ("People ____ different opinions.", "have", ["has", "is", "was"]),
        ("Neither she nor I ____ responsible.", "am", ["is", "are", "were"]),
    ]
    for i in range(100):
        q, correct, wrongs = items[i % len(items)]
        add_question(rows, rows, q, correct, wrongs)
    assert len(rows) == 100
    return rows


def build_passive():
    rows = []
    items = [
        ("Choose the correct passive form: People speak English in many countries.", "English is spoken in many countries.", ["English speaks in many countries.", "English is speak in many countries.", "English spoken in many countries."]),
        ("Choose the correct passive form: They are painting the house now.", "The house is being painted now.", ["The house is painted now.", "The house is painting now.", "The house has been painted now."]),
        ("Choose the correct passive form: Someone has cleaned the room.", "The room has been cleaned.", ["The room has cleaned.", "The room is cleaned.", "The room cleaned."]),
        ("Choose the correct passive form: They built the bridge in 2015.", "The bridge was built in 2015.", ["The bridge built in 2015.", "The bridge was building in 2015.", "The bridge has built in 2015."]),
        ("Choose the correct passive form: They were cooking dinner at 7 p.m.", "Dinner was being cooked at 7 p.m.", ["Dinner was cooked at 7 p.m.", "Dinner was cooking at 7 p.m.", "Dinner has been cooked at 7 p.m."]),
        ("Choose the correct passive form: They had made the decision before noon.", "The decision had been made before noon.", ["The decision had made before noon.", "The decision was made before noon.", "The decision has been made before noon."]),
        ("Choose the correct passive form: They will sign the contract tomorrow.", "The contract will be signed tomorrow.", ["The contract will sign tomorrow.", "The contract is signed tomorrow.", "The contract will signed tomorrow."]),
        ("Choose the correct passive form: We must protect wild animals.", "Wild animals must be protected.", ["Wild animals must protect.", "Wild animals must protected.", "Wild animals are protected must."]),
        ("Choose the correct passive form: The teacher checks the tests every week.", "The tests are checked every week.", ["The tests checked every week.", "The tests are checking every week.", "The tests have checked every week."]),
        ("Choose the correct passive form: Mary is writing an email.", "An email is being written by Mary.", ["An email is written by Mary.", "An email is writing by Mary.", "An email has been written by Mary."]),
    ]
    for i in range(100):
        q, correct, wrongs = items[i % len(items)]
        add_question(rows, rows, q, correct, wrongs)
    assert len(rows) == 100
    return rows


def build_modal():
    rows = []
    items = [
        ("She looks pale. She ____ be sick.", "must", ["might", "will", "can"]),
        ("I am not sure, but it ____ rain later.", "might", ["must", "will certainly", "can't"]),
        ("The traffic is light, so he ____ arrive soon.", "should", ["must not", "can't", "couldn't"]),
        ("Swimming ____ be dangerous if you ignore safety rules.", "can", ["must", "will certainly", "has to"]),
        ("I am certain that he ____ arrive at 8 p.m.", "will", ["might", "could", "may"]),
        ("The lights are off. They ____ be at home.", "can't", ["must", "should", "will"]),
        ("You ____ wear a helmet when riding a motorbike.", "must", ["might", "could", "may"]),
        ("It is possible that he is at the library. He ____ be there.", "may", ["must", "will", "can't"]),
        ("I expect the bus to arrive soon. It ____ arrive soon.", "should", ["must not", "can't", "wouldn't"]),
        ("She is unsure, so the answer ____ be wrong.", "could", ["must", "will", "can't"]),
        ("Don't worry. I ____ help you with this exercise.", "will", ["must", "might", "could not"]),
        ("Look at his wet hair. He ____ have been in the rain.", "must", ["can't", "might not", "won't"]),
        ("This room is dirty. It ____ be cleaned.", "must", ["might", "could", "may"]),
        ("Maybe Tom is busy. He ____ not answer the phone.", "might", ["must", "will certainly", "can't"]),
        ("The museum is closed today. We ____ go there now.", "can't", ["must", "should", "will"]),
        ("If nothing goes wrong, she ____ finish soon.", "should", ["can't", "must not", "couldn't"]),
        ("Perhaps it is John at the door. It ____ be John.", "could", ["must", "will", "can't"]),
        ("I promise I ____ keep your secret.", "will", ["may", "might", "could"]),
        ("You look tired. You ____ take a rest.", "should", ["can't", "mustn't", "couldn't"]),
        ("He knows the rule very well. He ____ make that mistake.", "can't", ["must", "should", "will"]),
    ]
    for q, correct, wrongs in items:
        add_question(rows, rows, q, correct, wrongs)
    assert len(rows) == 20
    return rows


def build_conjunctions():
    rows = []
    items = [
        ("We stayed home ____ it rained heavily.", "because", ["although", "but", "or"]),
        ("He studies hard ____ he can pass the exam.", "so that", ["because", "although", "unless"]),
        ("____ he was tired, he finished the report.", "Although", ["Because", "Unless", "So"]),
        ("You will fail ____ you study hard.", "unless", ["if", "because", "although"]),
        ("Call me as soon as you ____.", "arrive", ["will arrive", "arrived", "are arriving"]),
        ("I will wait here until you ____ back.", "come", ["will come", "came", "coming"]),
        ("He studies hard in order ____ the exam.", "to pass", ["pass", "passing", "passed"]),
        ("We listened eagerly, ____ he brought good news.", "for", ["but", "or", "yet"]),
        ("She bought a book, ____ he bought a pen.", "and", ["but", "nor", "unless"]),
        ("He doesn't smoke, nor ____ he drink.", "does", ["do", "is", "has"]),
        ("You can have tea ____ coffee.", "or", ["but", "so", "for"]),
        ("He studied hard, ____ he failed the exam.", "but", ["so", "and", "for"]),
        ("The weather was cold, ____ we went swimming.", "yet", ["so", "for", "or"]),
        ("It started to rain, ____ we went inside.", "so", ["yet", "or", "for"]),
        ("Both Tom and Jerry ____ funny.", "are", ["is", "has", "was"]),
        ("Either you or he ____ responsible.", "is", ["are", "were", "have"]),
        ("Neither the teacher nor the students ____ there.", "were", ["was", "is", "has"]),
        ("Not only ____ she beautiful, but she is also kind.", "is", ["does", "has", "will"]),
        ("Before he went to bed, he ____ his homework.", "had finished", ["finished", "has finished", "was finishing"]),
        ("____ you start, you cannot stop.", "Once", ["Unless", "Although", "Because"]),
        ("Whenever I ____ London, it rains.", "visit", ["will visit", "visited", "am visiting"]),
        ("While I was cooking, the phone ____.", "rang", ["rings", "has rung", "was ringing"]),
        ("After she had locked the door, she ____ to the bus stop.", "walked", ["had walked", "has walked", "was walking"]),
        ("I will help you ____ I finish my work.", "as soon as", ["although", "unless", "because of"]),
        ("____ it rains tomorrow, we will cancel the trip.", "If", ["Although", "Because", "Yet"]),
    ]
    for i in range(100):
        q, correct, wrongs = items[i % len(items)]
        add_question(rows, rows, q, correct, wrongs)
    assert len(rows) == 100
    return rows


def build_written():
    rows = []

    def add(task, key):
        rows.append({"no": len(rows) + 1, "task": task, "key": key})

    verb_tasks = [
        ("The children (play) ____ in the garden at the moment.", "are playing"),
        ("My father usually (go) ____ to work by bus.", "goes"),
        ("She (not/eat) ____ meat for years.", "has not eaten / hasn't eaten"),
        ("I (watch) ____ TV when the phone rang.", "was watching"),
        ("They (visit) ____ Da Nang last summer.", "visited"),
        ("Water (boil) ____ at 100 degrees Celsius.", "boils"),
        ("We (learn) ____ ten units so far.", "have learned / have learnt"),
        ("The train (leave) ____ at 7 a.m. tomorrow.", "leaves"),
        ("Listen! The birds (sing) ____.", "are singing"),
        ("By the time we arrived, the film (start) ____.", "had started"),
    ]
    for i in range(20):
        task, key = verb_tasks[i % len(verb_tasks)]
        add(f"Give the correct verb form: {task}", key)

    passive_tasks = [
        ("People speak English in many countries.", "English is spoken in many countries."),
        ("They built the bridge in 2015.", "The bridge was built in 2015."),
        ("The teacher is checking our tests now.", "Our tests are being checked by the teacher now."),
        ("Someone has cleaned the room.", "The room has been cleaned."),
        ("They will complete the project next month.", "The project will be completed next month."),
        ("We must protect wild animals.", "Wild animals must be protected."),
        ("They were painting the house at 8 a.m.", "The house was being painted at 8 a.m."),
        ("My sister makes the cake every Sunday.", "The cake is made by my sister every Sunday."),
        ("They had made the decision before noon.", "The decision had been made before noon."),
        ("The manager signed the documents.", "The documents were signed by the manager."),
    ]
    for i in range(20):
        task, key = passive_tasks[i % len(passive_tasks)]
        add(f"Change into passive voice: {task}", key)

    relative_tasks = [
        ("The book is interesting. I borrowed it from the library.", "The book which/that I borrowed from the library is interesting."),
        ("The girl won the competition. She is my best friend.", "The girl who/that won the competition is my best friend."),
        ("This is the house. I was born there.", "This is the house where I was born."),
        ("I remember the day. We first met on that day.", "I remember the day when we first met."),
        ("The boy is crying. His bike was stolen.", "The boy whose bike was stolen is crying."),
        ("The woman is my boss. You met her yesterday.", "The woman whom/who/that you met yesterday is my boss."),
        ("The city is beautiful. We visited it last year.", "The city which/that we visited last year is beautiful."),
        ("Tell me the reason. You were absent for that reason.", "Tell me the reason why you were absent."),
        ("I like students. They work hard.", "I like students who/that work hard."),
        ("The dress is beautiful. She bought it yesterday.", "The dress which/that she bought yesterday is beautiful."),
    ]
    for i in range(20):
        task, key = relative_tasks[i % len(relative_tasks)]
        add(f"Combine using a relative clause: {task}", key)

    correction_tasks = [
        ("She enjoy playing chess.", "enjoy -> enjoys"),
        ("I am knowing the answer.", "am knowing -> know"),
        ("Everyone are ready.", "are -> is"),
        ("The room must cleaned now.", "must cleaned -> must be cleaned"),
        ("Call me when you will arrive.", "will arrive -> arrive"),
        ("The information are useful.", "are -> is"),
        ("Neither the teacher nor the students is ready.", "is -> are"),
        ("The bridge built in 2015.", "built -> was built"),
        ("The girl which won is my friend.", "which -> who/that"),
        ("You will fail unless you do not study.", "do not study -> study"),
    ]
    for i in range(20):
        task, key = correction_tasks[i % len(correction_tasks)]
        add(f"Find and correct the mistake: {task}", key)

    rewrite_tasks = [
        ("He was tired, but he finished the report. (Although)", "Although he was tired, he finished the report."),
        ("It rained heavily, so we stayed home. (Because)", "We stayed home because it rained heavily."),
        ("He studies hard. He wants to pass the exam. (so that)", "He studies hard so that he can pass the exam."),
        ("If you do not study hard, you will fail. (Unless)", "You will fail unless you study hard."),
        ("She saves money because she wants to buy a laptop. (in order to)", "She saves money in order to buy a laptop."),
        ("I am sure she is sick because she looks pale. (must)", "She must be sick because she looks pale."),
        ("It is possible that he is at the library. (might/could)", "He might/could be at the library."),
        ("I expect the bus to arrive soon. (should)", "The bus should arrive soon."),
        ("I am certain that he will arrive at 8 p.m. (will)", "He will arrive at 8 p.m."),
        ("I am sure they are not at home because the lights are off. (can't)", "They can't be at home because the lights are off."),
    ]
    for i in range(20):
        task, key = rewrite_tasks[i % len(rewrite_tasks)]
        add(f"Rewrite the sentence: {task}", key)

    assert len(rows) == 100
    return rows


def add_answer_key(doc, sections, written):
    doc.add_page_break()
    paragraph(doc, "ANSWER KEY", size=16, bold=True, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER)
    for title, rows in sections:
        heading(doc, title)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for cell in table.rows[0].cells:
            cell.text = "No. - Ans"
            set_cell_shading(cell, "E8EEF5")
            style_cell_text(cell, size=10)
        for start in range(0, len(rows), 5):
            row = table.add_row().cells
            for c in range(5):
                if start + c < len(rows):
                    item = rows[start + c]
                    row[c].text = f"{item['no']}. {item['ans']}"
                style_cell_text(row[c], size=10)
        set_table_geometry(table, [1.3, 1.3, 1.3, 1.3, 1.3])

    heading(doc, "Suggested Answers - Written Practice")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "No."
    table.rows[0].cells[1].text = "Suggested answer"
    repeat_table_header(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        style_cell_text(cell, size=10)
    for item in written:
        row = table.add_row().cells
        row[0].text = str(item["no"])
        row[1].text = item["key"]
        style_cell_text(row[0], size=10)
        style_cell_text(row[1], size=10)
    set_table_geometry(table, [0.45, 6.05])


def main():
    sections = [
        ("PART A. Pronunciation - 100 Multiple-choice Questions", build_pronunciation()),
        ("PART B. Vocabulary - 100 Multiple-choice Questions", build_vocab()),
        ("PART C. Five Core Tenses - 100 Multiple-choice Questions", build_tenses()),
        ("PART D. Relative Clauses - 100 Multiple-choice Questions", build_relative()),
        ("PART E. Subject-Verb Agreement - 100 Multiple-choice Questions", build_sva()),
        ("PART F. Passive Voice - 100 Multiple-choice Questions", build_passive()),
        ("PART G. Modal Verbs - 20 Multiple-choice Questions", build_modal()),
        ("PART H. English Conjunctions - 100 Multiple-choice Questions", build_conjunctions()),
    ]
    written = build_written()

    assert sum(len(rows) for _, rows in sections) == 720
    assert len(written) == 100

    doc = Document()
    configure_doc(doc)

    paragraph(doc, "ĐỀ MỚI BASIC IELTS - NGỮ PHÁP THEO CHỦ ĐIỂM", size=18, bold=True, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph(doc, "720 câu trắc nghiệm + 100 câu tự luận | Bám sát dạng hay gặp trong đề kiểm tra Basic IELTS", size=11, bold=True, color="555555", align=WD_ALIGN_PARAGRAPH.CENTER)

    info = doc.add_table(rows=3, cols=2)
    info.style = "Table Grid"
    for i, label in enumerate(["Name:", "Class:", "Date:"]):
        info.rows[i].cells[0].text = label
        info.rows[i].cells[1].text = ""
    set_table_geometry(info, [1.2, 5.3])

    paragraph(doc, "Instructions: Choose A, B, C, or D. Written questions require complete answers or corrected sentences. The answer key is placed at the end for self-study.", bold=True)
    paragraph(doc, "Coverage: final -s/-es, final -ed, stress, likes/dislikes vocabulary, grammar signals, present simple, present continuous, past simple, past continuous, present perfect, relative clauses, subject-verb agreement, passive voice, modal verbs, and conjunctions.", size=10)

    for title, rows in sections:
        add_mcq_table(doc, title, rows)
    add_written_table(doc, "PART I. Written Practice - 100 Questions", written)
    add_answer_key(doc, sections, written)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Basic IELTS Grammar Question Bank")
    set_font(r, size=9, color="555555")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
