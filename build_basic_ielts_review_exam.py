from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "De_on_tap_Basic_IELTS_ngu_phap_250_cau.docx"


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_col_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_col_width(cell, Inches(widths[i]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph(doc, text="", style=None, bold=False, size=11, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_section_heading(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_font(r, size=13, bold=True, color="2E74B5")


def add_question_table(doc, rows, title):
    add_section_heading(doc, title)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0]
    hdr.cells[0].text = "No."
    hdr.cells[1].text = "Question"
    repeat_table_header(hdr)
    for cell in hdr.cells:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, bold=True)
    for no, question, options, answer in rows:
        row = table.add_row()
        row.cells[0].text = str(no)
        text = question + "\n" + "\n".join(
            f"{chr(65 + i)}. {opt}" for i, opt in enumerate(options)
        )
        row.cells[1].text = text
        for p in row.cells[1].paragraphs:
            p.paragraph_format.space_after = Pt(0)
    set_table_geometry(table, [0.45, 6.05])
    return table


def add_short_answer_table(doc, rows, title):
    add_section_heading(doc, title)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0]
    hdr.cells[0].text = "No."
    hdr.cells[1].text = "Task"
    hdr.cells[2].text = "Answer"
    repeat_table_header(hdr)
    for cell in hdr.cells:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                set_font(r, bold=True)
    for no, task, key in rows:
        row = table.add_row()
        row.cells[0].text = str(no)
        row.cells[1].text = task
        row.cells[2].text = "\n\n"
    set_table_geometry(table, [0.45, 4.75, 1.30])
    return table


def mc(no, question, options, answer):
    return (no, question, options, answer)


def sa(no, task, key):
    return (no, task, key)


mcq = []

pronunciation = [
    ("Choose the word whose final -s/-es is pronounced differently.", ["books", "cats", "maps", "dogs"], "D"),
    ("Choose the word whose final -s/-es is pronounced differently.", ["buses", "watches", "boxes", "plays"], "D"),
    ("Choose the word whose final -s/-es is pronounced differently.", ["laughs", "cooks", "hopes", "lives"], "D"),
    ("Choose the word whose final -s/-es is pronounced differently.", ["villages", "changes", "judges", "desks"], "D"),
    ("Choose the word whose final -s/-es is pronounced differently.", ["pens", "rooms", "songs", "cups"], "D"),
    ("Choose the word whose final -s/-es is pronounced differently.", ["classes", "dishes", "kisses", "schools"], "D"),
    ("Choose the word whose final -s/-es is pronounced differently.", ["takes", "makes", "stops", "sees"], "D"),
    ("Choose the word whose final -s/-es is pronounced differently.", ["months", "hats", "cliffs", "bags"], "D"),
    ("Choose the word whose final -s/-es is pronounced differently.", ["houses", "freezes", "dances", "keys"], "D"),
    ("Choose the word whose final -s/-es is pronounced differently.", ["students", "desks", "cups", "islands"], "D"),
    ("Which word has the final sound /iz/?", ["boxes", "books", "dreams", "keys"], "A"),
    ("Which word has the final sound /iz/?", ["plays", "watches", "rooms", "bags"], "B"),
    ("Which word has the final sound /iz/?", ["jobs", "pens", "wishes", "schools"], "C"),
    ("Which word has the final sound /iz/?", ["maps", "hopes", "students", "bridges"], "D"),
    ("Which word has the final sound /s/?", ["coughs", "goes", "lives", "plans"], "A"),
    ("Which word has the final sound /s/?", ["songs", "takes", "keys", "islands"], "B"),
    ("Which word has the final sound /z/?", ["cups", "cliffs", "dreams", "months"], "C"),
    ("Which word has the final sound /z/?", ["laughs", "desks", "hats", "belongs"], "D"),
    ("Which word has the final sound /iz/?", ["services", "students", "maps", "rooms"], "A"),
    ("Which word has the final sound /z/?", ["stops", "types", "beliefs", "drives"], "D"),
    ("Choose the word whose final -ed is pronounced differently.", ["wanted", "needed", "visited", "played"], "D"),
    ("Choose the word whose final -ed is pronounced differently.", ["watched", "cooked", "stopped", "opened"], "D"),
    ("Choose the word whose final -ed is pronounced differently.", ["called", "cleaned", "lived", "laughed"], "D"),
    ("Choose the word whose final -ed is pronounced differently.", ["decided", "started", "helped", "painted"], "C"),
    ("Choose the word whose final -ed is pronounced differently.", ["washed", "missed", "hoped", "studied"], "D"),
    ("Choose the word whose final -ed is pronounced differently.", ["changed", "enjoyed", "planned", "walked"], "D"),
    ("Choose the word whose final -ed is pronounced differently.", ["waited", "asked", "added", "ended"], "B"),
    ("Choose the word whose final -ed is pronounced differently.", ["looked", "fixed", "danced", "offered"], "D"),
    ("Choose the word whose final -ed is pronounced differently.", ["arrived", "loved", "used", "wanted"], "D"),
    ("Choose the word whose final -ed is pronounced differently.", ["packed", "checked", "stopped", "borrowed"], "D"),
    ("Choose the word with a different stress pattern.", ["artist", "teacher", "doctor", "enjoy"], "D"),
    ("Choose the word with a different stress pattern.", ["begin", "prefer", "happen", "agree"], "C"),
    ("Choose the word with a different stress pattern.", ["holiday", "family", "important", "animal"], "C"),
    ("Choose the word with a different stress pattern.", ["invite", "visit", "finish", "open"], "A"),
    ("Choose the word with a different stress pattern.", ["attention", "expensive", "beautiful", "computer"], "C"),
    ("Choose the word with a different stress pattern.", ["museum", "musician", "vacation", "camera"], "D"),
    ("Choose the word with a different stress pattern.", ["yesterday", "banana", "tomato", "potato"], "A"),
    ("Choose the word with a different stress pattern.", ["engineer", "volunteer", "magazine", "newspaper"], "D"),
    ("Choose the word with a different stress pattern.", ["afternoon", "understand", "exercise", "Japanese"], "C"),
    ("Choose the word with a different stress pattern.", ["correct", "polite", "simple", "arrive"], "C"),
]

vocab = [
    ("I really enjoy ____ English podcasts in the evening.", ["listen", "listening", "to listening", "listened"], "B"),
    ("She prefers tea ____ coffee.", ["than", "from", "to", "with"], "C"),
    ("Tom is keen ____ playing chess after school.", ["on", "in", "at", "for"], "A"),
    ("I am not a big fan ____ spicy food.", ["of", "with", "about", "on"], "A"),
    ("My brother has a passion ____ vintage cameras.", ["at", "for", "with", "to"], "B"),
    ("They can't stand ____ in long queues.", ["wait", "to waiting", "waiting", "waited"], "C"),
    ("The word closest in meaning to 'detest' is ____.", ["like", "hate strongly", "prefer", "enjoy"], "B"),
    ("A person who is 'fond of' music ____ music.", ["dislikes", "knows nothing about", "likes", "avoids"], "C"),
    ("Choose the best expression: 'I ____ horror films; they make me nervous.'", ["am into", "can't bear", "am fond of", "prefer"], "B"),
    ("'Be into something' means ____.", ["be interested in it", "be late for it", "be afraid of it", "be responsible for it"], "A"),
    ("A 'service' is ____.", ["a kind of meal only", "help or work provided for others", "a school subject", "a final sound"], "B"),
    ("A 'choice' is ____.", ["a decision between options", "a place to live", "a kind of illness", "a daily habit"], "A"),
    ("A 'sentence' is ____.", ["a cooking tool", "a group of words with meaning", "a plan for travel", "a sport"], "B"),
    ("'Recently' is closest in meaning to ____.", ["a long time ago", "many years later", "not long ago", "never"], "C"),
    ("'So far' means ____.", ["until now", "tomorrow morning", "in the past only", "at once"], "A"),
    ("'At the moment' is a signal phrase for ____.", ["present continuous", "past perfect", "future simple only", "passive voice only"], "A"),
    ("'Every day' usually signals ____.", ["present simple", "present perfect", "past continuous", "modal deduction"], "A"),
    ("'Yesterday' usually signals ____.", ["present continuous", "past simple", "will", "present perfect"], "B"),
    ("'Look at those dark clouds!' usually gives evidence for ____.", ["will", "be going to", "present perfect", "past simple"], "B"),
    ("'I think...' often introduces a future ____ with will.", ["certainty based on proof", "subjective prediction", "past habit", "passive action"], "B"),
    ("The word 'advice' is ____.", ["countable plural", "uncountable", "always plural", "a verb"], "B"),
    ("The word 'people' usually takes a ____ verb.", ["singular", "plural", "passive only", "modal only"], "B"),
    ("'The rich' refers to ____.", ["one rich person", "rich people in general", "money only", "a place"], "B"),
    ("'Either...or' expresses ____.", ["addition", "choice", "reason", "time"], "B"),
    ("'Although' introduces ____.", ["contrast/concession", "purpose only", "ownership", "certainty"], "A"),
    ("'So that' introduces ____.", ["purpose", "contrast", "possession", "pronunciation"], "A"),
    ("'Unless' means ____.", ["if not", "because", "although", "as soon as"], "A"),
    ("The pronoun 'whose' shows ____.", ["place", "time", "possession", "reason"], "C"),
    ("The relative adverb 'where' refers to ____.", ["a person", "a place", "a reason", "a possession"], "B"),
    ("The modal 'must' often shows ____.", ["weak possibility", "strong logical deduction", "past habit", "dislike"], "B"),
]

tenses = [
    ("My parents ____ up early every morning.", ["wake", "wakes", "are waking", "woke"], "A"),
    ("Listen! The birds ____ outside.", ["sing", "sings", "are singing", "sang"], "C"),
    ("Water ____ at 100 degrees Celsius.", ["boil", "boils", "is boiling", "boiled"], "B"),
    ("She ____ meat because she is a vegetarian.", ["doesn't eat", "isn't eating", "didn't eat", "hasn't eating"], "A"),
    ("We ____ for the bus at the moment.", ["wait", "waited", "are waiting", "have waited"], "C"),
    ("He is always ____ late for class.", ["come", "comes", "coming", "came"], "C"),
    ("I ____ what you mean.", ["am knowing", "know", "knew", "have knowing"], "B"),
    ("The train ____ at 8 a.m. tomorrow.", ["leaves", "is leaving", "left", "has left"], "A"),
    ("She ____ to London next Friday. She has bought the ticket.", ["flies", "is flying", "flew", "has flown"], "B"),
    ("They usually ____ football after school.", ["play", "plays", "are playing", "played"], "A"),
    ("I ____ my grandmother last weekend.", ["visit", "visited", "have visited", "was visiting"], "B"),
    ("When I got home, my sister ____ dinner.", ["cooked", "was cooking", "has cooked", "is cooking"], "B"),
    ("At 9 p.m. yesterday, we ____ a movie.", ["watched", "watch", "were watching", "have watched"], "C"),
    ("She stood up, ____ her bag, and left.", ["takes", "took", "was taking", "has taken"], "B"),
    ("While my dad ____ the car, I was cleaning my room.", ["washed", "was washing", "has washed", "washes"], "B"),
    ("The phone ____ while I was doing homework.", ["rang", "was ringing", "rings", "has rung"], "A"),
    ("They ____ in this town in 2010.", ["live", "lived", "have lived", "were living"], "B"),
    ("I ____ when the alarm rang.", ["slept", "was sleeping", "sleep", "have slept"], "B"),
    ("My friends ____ to the museum two days ago.", ["go", "went", "have gone", "were going"], "B"),
    ("It ____ heavily when we left the office.", ["rained", "was raining", "has rained", "rains"], "B"),
    ("Mary ____ any meat for years.", ["doesn't eat", "didn't eat", "hasn't eaten", "isn't eating"], "C"),
    ("I ____ this film twice.", ["see", "saw", "have seen", "am seeing"], "C"),
    ("Have you ____ finished your homework?", ["yet", "already", "ever", "since"], "B"),
    ("We haven't met him ____.", ["already", "yet", "ever", "just"], "B"),
    ("She has ____ visited Japan before.", ["ever", "never", "ago", "last"], "B"),
    ("They have lived here ____ 2020.", ["for", "since", "during", "ago"], "B"),
    ("He has studied English ____ five years.", ["since", "for", "ago", "last"], "B"),
    ("Over the past month, prices ____ quickly.", ["increase", "increased", "have increased", "were increasing"], "C"),
    ("I have ____ eaten lunch, so I am not hungry.", ["just", "yet", "ago", "last"], "A"),
    ("____ you ever tried Indian food?", ["Do", "Did", "Have", "Are"], "C"),
    ("Oh no, I forgot my pen. I ____ lend you one.", ["am going to", "will", "am", "have"], "B"),
    ("Look at those clouds! It ____ rain.", ["will", "is going to", "has", "was"], "B"),
    ("I think she ____ pass the exam.", ["will", "is going to", "has", "was"], "A"),
    ("We ____ visit Paris this summer. We booked the tickets.", ["will", "are going to", "visited", "visit"], "B"),
    ("Don't worry. I ____ keep your secret.", ["am going to", "will", "was", "have"], "B"),
    ("Watch out! The glass ____ fall.", ["will", "is going to", "has", "did"], "B"),
    ("Maybe they ____ arrive late.", ["are going to", "will", "were", "have"], "B"),
    ("She ____ buy a new laptop. She has saved money for months.", ["will", "is going to", "has", "does"], "B"),
    ("I am sure he ____ be here soon.", ["will", "is going to", "was", "has"], "A"),
    ("The baby is very tired. She ____ fall asleep.", ["will", "is going to", "has", "did"], "B"),
    ("By the time we arrived, the film ____ already started.", ["has", "had", "was", "is"], "B"),
    ("After he ____ his homework, he went to bed.", ["finished", "had finished", "has finished", "was finishing"], "B"),
    ("Before she moved to Hanoi, she ____ in Hue.", ["lived", "has lived", "had lived", "is living"], "C"),
    ("When the teacher entered, the students ____ loudly.", ["talked", "were talking", "have talked", "talk"], "B"),
    ("I ____ my keys. I can't open the door now.", ["lost", "was losing", "have lost", "had lost"], "C"),
    ("She ____ in the garden now.", ["works", "worked", "is working", "has worked"], "C"),
    ("My brother rarely ____ breakfast.", ["skip", "skips", "is skipping", "skipped"], "B"),
    ("They ____ TV when the lights went out.", ["watched", "were watching", "watch", "have watched"], "B"),
    ("So far, we ____ ten grammar units.", ["finish", "finished", "have finished", "were finishing"], "C"),
    ("The shop ____ at 9 p.m. every day.", ["closes", "is closing", "closed", "has closed"], "A"),
]

grammar_mid = [
    ("Everyone ____ to be happy.", ["want", "wants", "are wanting", "have wanted"], "B"),
    ("The information ____ useful.", ["is", "are", "were", "be"], "A"),
    ("Ten dollars ____ not enough for lunch.", ["is", "are", "were", "have"], "A"),
    ("Tom and Jerry ____ famous cartoon characters.", ["is", "are", "was", "be"], "B"),
    ("The police ____ investigating the case.", ["is", "are", "has", "was"], "B"),
    ("A number of students ____ absent today.", ["is", "are", "has", "was"], "B"),
    ("The number of students in this class ____ thirty.", ["is", "are", "were", "have"], "A"),
    ("Each student ____ a dictionary.", ["have", "has", "are having", "were having"], "B"),
    ("Swimming ____ good for your health.", ["is", "are", "were", "have"], "A"),
    ("Neither the teacher nor the students ____ ready.", ["is", "are", "has", "was"], "B"),
    ("Either my parents or my brother ____ me at the station.", ["meet", "meets", "are meeting", "were meeting"], "B"),
    ("A lot of water ____ wasted every day.", ["is", "are", "were", "have"], "A"),
    ("Some of the animals ____ in danger.", ["is", "are", "has", "was"], "B"),
    ("My mother, as well as my sisters, ____ at home.", ["is", "are", "were", "have"], "A"),
    ("The rich ____ not always happy.", ["is", "are", "has", "was"], "B"),
    ("She looks very tired. She ____ be sick.", ["must", "can", "will", "would"], "A"),
    ("It is possible that he will come late. He ____ come late.", ["must", "might", "should", "will certainly"], "B"),
    ("The traffic is light, so he ____ arrive soon.", ["should", "must not", "could not", "can't"], "A"),
    ("I am certain that the train ____ arrive at 7.", ["might", "could", "will", "may"], "C"),
    ("Swimming ____ be dangerous if you ignore safety rules.", ["can", "must", "will certainly", "has to"], "A"),
    ("She is not sure, but it ____ be the right answer.", ["must", "might", "will", "ought"], "B"),
    ("You ____ wear a helmet when riding a motorbike.", ["must", "might", "could", "may"], "A"),
    ("He has studied hard; he ____ pass the test.", ["should", "must not", "can't", "couldn't"], "A"),
    ("This room ____ be cleaned before guests arrive.", ["must", "might", "could", "may"], "A"),
    ("Look! The lights are off. They ____ be at home.", ["must", "can't", "should", "will"], "B"),
    ("The letter ____ every morning.", ["delivers", "is delivered", "delivered", "has delivered"], "B"),
    ("The house ____ at the moment.", ["paints", "is painting", "is being painted", "painted"], "C"),
    ("The homework ____ already ____.", ["has / done", "has / been done", "is / done", "was / doing"], "B"),
    ("The car ____ yesterday.", ["repaired", "was repaired", "is repairing", "has repaired"], "B"),
    ("Dinner ____ when I came home.", ["was being cooked", "cooked", "is cooking", "has cooked"], "A"),
    ("The decision ____ before the meeting started.", ["made", "had made", "had been made", "has making"], "C"),
    ("The contract ____ next week.", ["will sign", "will be signed", "is signing", "signed"], "B"),
    ("This problem ____ by the students.", ["can solve", "can be solved", "can solved", "can be solve"], "B"),
    ("English ____ in many countries.", ["speaks", "is spoken", "is speaking", "has spoken"], "B"),
    ("The new bridge ____ in 2015.", ["built", "was built", "has built", "was building"], "B"),
    ("The report ____ now by the manager.", ["is checked", "is being checked", "checks", "has checked"], "B"),
    ("The window ____ by the boys last night.", ["broke", "was broken", "has broken", "is breaking"], "B"),
    ("The exercises ____ by next Friday.", ["will finish", "will be finished", "will be finishing", "finish"], "B"),
    ("These books ____ carefully.", ["must keep", "must be kept", "must kept", "are keeping"], "B"),
    ("The cake ____ by my sister every Sunday.", ["makes", "is made", "made", "has made"], "B"),
]

clauses = [
    ("The man ____ is standing there is my uncle.", ["who", "which", "where", "when"], "A"),
    ("The book ____ I borrowed is very useful.", ["who", "which", "where", "why"], "B"),
    ("The girl ____ hair is long is my friend.", ["who", "whose", "which", "where"], "B"),
    ("This is the house ____ I was born.", ["who", "which", "where", "why"], "C"),
    ("Do you remember the day ____ we first met?", ["when", "where", "why", "whose"], "A"),
    ("Tell me the reason ____ you are late.", ["where", "when", "why", "whose"], "C"),
    ("The woman ____ you met yesterday is my boss.", ["whom", "where", "when", "whose"], "A"),
    ("The car ____ he bought is expensive.", ["where", "that", "when", "why"], "B"),
    ("I like people ____ are honest.", ["which", "who", "where", "when"], "B"),
    ("The city ____ we visited was beautiful.", ["who", "whose", "which", "why"], "C"),
    ("We stayed home ____ it rained heavily.", ["because", "although", "but", "or"], "A"),
    ("He studies hard ____ he can pass the exam.", ["so that", "because", "although", "unless"], "A"),
    ("She went to bed early ____ she was tired.", ["because", "so that", "unless", "or"], "A"),
    ("____ he was tired, he finished the report.", ["Although", "Because", "Unless", "So"], "A"),
    ("You will fail ____ you study hard.", ["if", "unless", "because", "although"], "B"),
    ("Call me as soon as you ____.", ["will arrive", "arrive", "arrived", "are arriving"], "B"),
    ("I will wait here until you ____ back.", ["will come", "came", "come", "coming"], "C"),
    ("When I ____ home tomorrow, I will call you.", ["get", "will get", "got", "am getting"], "A"),
    ("He studies hard in order ____ the exam.", ["pass", "to pass", "passing", "passed"], "B"),
    ("We listened eagerly, ____ he brought good news.", ["for", "but", "or", "yet"], "A"),
    ("She bought a book, ____ he bought a pen.", ["and", "but", "nor", "unless"], "A"),
    ("He doesn't smoke, nor ____ he drink.", ["do", "does", "is", "has"], "B"),
    ("You can have tea ____ coffee.", ["but", "so", "or", "for"], "C"),
    ("He studied hard, ____ he failed the exam.", ["but", "so", "and", "for"], "A"),
    ("It started to rain, ____ we went inside.", ["so", "yet", "or", "for"], "A"),
    ("Both Tom and Jerry ____ funny.", ["is", "are", "has", "was"], "B"),
    ("Either you or he ____ responsible for this.", ["are", "is", "were", "have"], "B"),
    ("Neither the teacher nor the students ____ there.", ["was", "is", "were", "has"], "C"),
    ("Not only ____ she beautiful, but she is also kind.", ["is", "does", "has", "will"], "A"),
    ("Before he went to bed, he ____ his homework.", ["finished", "has finished", "had finished", "was finishing"], "C"),
    ("____ you start, you cannot stop.", ["Once", "Unless", "Although", "Because"], "A"),
    ("Whenever I ____ London, it rains.", ["visit", "will visit", "visited", "am visiting"], "A"),
    ("While I was cooking, the phone ____.", ["rings", "rang", "has rung", "was ringing"], "B"),
    ("After she had locked the door, she ____ to the bus stop.", ["walked", "had walked", "has walked", "was walking"], "A"),
    ("I will help you ____ I finish my work.", ["as soon as", "although", "unless", "because of"], "A"),
    ("____ it rains tomorrow, we will cancel the trip.", ["If", "Although", "Because", "Yet"], "A"),
    ("She is the student ____ won the prize.", ["which", "who", "where", "why"], "B"),
    ("The village ____ my grandparents live is peaceful.", ["where", "which", "who", "when"], "A"),
    ("The year ____ we moved here was 2020.", ["where", "when", "why", "whose"], "B"),
    ("This is the boy ____ bike was stolen.", ["who", "which", "whose", "where"], "C"),
]

for group in [pronunciation, vocab, tenses, grammar_mid, clauses]:
    for question, options, answer in group:
        mcq.append(mc(len(mcq) + 1, question, options, answer))

assert len(mcq) == 200


short_answers = [
    sa(1, "Give the correct verb form: The children (play) __________ in the garden at the moment.", "are playing"),
    sa(2, "Give the correct verb form: My father usually (go) __________ to work by bus.", "goes"),
    sa(3, "Give the correct verb form: She (not/eat) __________ meat for years.", "has not eaten / hasn't eaten"),
    sa(4, "Give the correct verb form: I (watch) __________ TV when the phone rang.", "was watching"),
    sa(5, "Give the correct verb form: They (visit) __________ Da Nang last summer.", "visited"),
    sa(6, "Give the correct verb form: Look! The baby (fall) __________ asleep.", "is going to fall"),
    sa(7, "Give the correct verb form: Before he went out, he (finish) __________ his homework.", "had finished"),
    sa(8, "Give the correct verb form: Water (boil) __________ at 100 degrees Celsius.", "boils"),
    sa(9, "Give the correct verb form: We (learn) __________ ten units so far.", "have learned / have learnt"),
    sa(10, "Give the correct verb form: The train (leave) __________ at 7 a.m. tomorrow.", "leaves"),
    sa(11, "Change into passive voice: People speak English in many countries.", "English is spoken in many countries."),
    sa(12, "Change into passive voice: They built the bridge in 2015.", "The bridge was built in 2015."),
    sa(13, "Change into passive voice: The teacher is checking our tests now.", "Our tests are being checked by the teacher now."),
    sa(14, "Change into passive voice: Someone has cleaned the room.", "The room has been cleaned."),
    sa(15, "Change into passive voice: They will complete the project next month.", "The project will be completed next month."),
    sa(16, "Change into active voice: The documents were signed by the manager.", "The manager signed the documents."),
    sa(17, "Change into active voice: The homework has been done by Mary.", "Mary has done the homework."),
    sa(18, "Change into passive voice: We must protect wild animals.", "Wild animals must be protected."),
    sa(19, "Change into passive voice: They were painting the house at 8 a.m.", "The house was being painted at 8 a.m."),
    sa(20, "Change into active voice: The cake is made by my sister every Sunday.", "My sister makes the cake every Sunday."),
    sa(21, "Combine with a relative clause: The book is interesting. I borrowed it from the library.", "The book which/that I borrowed from the library is interesting."),
    sa(22, "Combine with a relative clause: The girl won the competition. She is my best friend.", "The girl who/that won the competition is my best friend."),
    sa(23, "Combine with a relative clause: This is the house. I was born there.", "This is the house where I was born."),
    sa(24, "Combine with a relative clause: I remember the day. We first met on that day.", "I remember the day when we first met."),
    sa(25, "Combine with a relative clause: The boy is crying. His bike was stolen.", "The boy whose bike was stolen is crying."),
    sa(26, "Combine with a relative clause: The woman is my boss. You met her yesterday.", "The woman whom/who/that you met yesterday is my boss."),
    sa(27, "Combine with a relative clause: The city is beautiful. We visited it last year.", "The city which/that we visited last year is beautiful."),
    sa(28, "Combine with a relative clause: Tell me the reason. You were absent for that reason.", "Tell me the reason why you were absent."),
    sa(29, "Combine with a relative clause: I like students. They work hard.", "I like students who/that work hard."),
    sa(30, "Combine with a relative clause: The dress is beautiful. She bought it yesterday.", "The dress which/that she bought yesterday is beautiful."),
    sa(31, "Rewrite using although: He was tired, but he finished the report.", "Although he was tired, he finished the report."),
    sa(32, "Rewrite using because: It rained heavily, so we stayed home.", "We stayed home because it rained heavily."),
    sa(33, "Rewrite using so that: He studies hard. He wants to pass the exam.", "He studies hard so that he can pass the exam."),
    sa(34, "Rewrite using unless: If you do not study hard, you will fail.", "You will fail unless you study hard."),
    sa(35, "Rewrite using in order to: She saves money because she wants to buy a laptop.", "She saves money in order to buy a laptop."),
    sa(36, "Rewrite with must: I am sure she is sick because she looks pale.", "She must be sick because she looks pale."),
    sa(37, "Rewrite with might/could: It is possible that he is at the library.", "He might/could be at the library."),
    sa(38, "Rewrite with should: I expect the bus to arrive soon.", "The bus should arrive soon."),
    sa(39, "Rewrite with will: I am certain that he will arrive at 8 p.m.", "He will arrive at 8 p.m."),
    sa(40, "Rewrite with can't: I am sure they are not at home because the lights are off.", "They can't be at home because the lights are off."),
    sa(41, "Find and correct the mistake: She enjoy playing chess.", "enjoy -> enjoys"),
    sa(42, "Find and correct the mistake: I am knowing the answer.", "am knowing -> know"),
    sa(43, "Find and correct the mistake: Everyone are ready.", "are -> is"),
    sa(44, "Find and correct the mistake: The room must cleaned now.", "must cleaned -> must be cleaned"),
    sa(45, "Find and correct the mistake: Call me when you will arrive.", "will arrive -> arrive"),
    sa(46, "Write 50-70 words about a hobby you enjoy. Use at least one expression of liking.", "Student's own answer."),
    sa(47, "Write 50-70 words about a past holiday. Use at least two past simple verbs.", "Student's own answer."),
    sa(48, "Write 50-70 words about your plan for next weekend. Use will or be going to correctly.", "Student's own answer."),
    sa(49, "Write 50-70 words about a person you admire. Use at least one relative clause.", "Student's own answer."),
    sa(50, "Write 50-70 words about how to study English better. Use at least one modal verb.", "Student's own answer."),
]

assert len(short_answers) == 50


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def build_doc():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ĐỀ ÔN TẬP BASIC IELTS - NGỮ PHÁP")
    set_font(run, size=18, bold=True, color="1F4D78")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("200 câu trắc nghiệm + 50 câu tự luận | Dựa trên tài liệu ngữ pháp ôn Basic IELTS")
    set_font(r, size=11, bold=True, color="555555")

    info = doc.add_table(rows=3, cols=2)
    info.style = "Table Grid"
    labels = ["Name:", "Class:", "Date:"]
    for i, label in enumerate(labels):
        info.rows[i].cells[0].text = label
        info.rows[i].cells[1].text = ""
    set_table_geometry(info, [1.2, 5.3])

    add_paragraph(
        doc,
        "Instructions: Choose A, B, C, or D for multiple-choice questions. For written questions, write full answers in the space provided. Suggested answers are at the end of this document.",
        bold=True,
    )

    add_question_table(doc, mcq[0:40], "PART A. Pronunciation and Stress (Questions 1-40)")
    add_question_table(doc, mcq[40:70], "PART B. Vocabulary and Grammar Signals (Questions 41-70)")
    add_question_table(doc, mcq[70:120], "PART C. Tenses and Future Forms (Questions 71-120)")
    add_question_table(doc, mcq[120:160], "PART D. Subject-Verb Agreement, Modals, Passive Voice (Questions 121-160)")
    add_question_table(doc, mcq[160:200], "PART E. Relative Clauses and Conjunctions (Questions 161-200)")
    add_short_answer_table(doc, short_answers, "PART F. Written Practice (Questions 1-50)")

    doc.add_page_break()
    add_paragraph(doc, "ANSWER KEY", style=None, bold=True, size=16, color="1F4D78", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_section_heading(doc, "Multiple-choice answers")
    ans_table = doc.add_table(rows=1, cols=5)
    ans_table.style = "Table Grid"
    for i, cell in enumerate(ans_table.rows[0].cells):
        cell.text = f"No. - Ans" if i == 0 else "No. - Ans"
        set_cell_shading(cell, "E8EEF5")
    for start in range(0, 200, 5):
        row = ans_table.add_row()
        for c in range(5):
            q = mcq[start + c]
            row.cells[c].text = f"{q[0]}. {q[3]}"
    set_table_geometry(ans_table, [1.3, 1.3, 1.3, 1.3, 1.3])

    add_section_heading(doc, "Suggested answers for written practice")
    key_table = doc.add_table(rows=1, cols=2)
    key_table.style = "Table Grid"
    key_table.rows[0].cells[0].text = "No."
    key_table.rows[0].cells[1].text = "Suggested answer"
    repeat_table_header(key_table.rows[0])
    for cell in key_table.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
    for no, task, key in short_answers:
        row = key_table.add_row()
        row.cells[0].text = str(no)
        row.cells[1].text = key
    set_table_geometry(key_table, [0.45, 6.05])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Basic IELTS Grammar Review")
    set_font(r, size=9, color="555555")

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
